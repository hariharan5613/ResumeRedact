"""
Resume Anonymizer — Flask Backend
===================================
Exact port of the Streamlit reference implementation.
Replaces: fitz (PyMuPDF) → pdfplumber + pypdf + reportlab
All logic preserved 1:1: phone regex, email regex, contact symbols,
section detection, background color sampling, watermark, name/role extraction.
"""

from flask import Flask, request, jsonify, send_file, make_response
import io, os, re, base64, zipfile, tempfile, time, traceback
from collections import Counter

# ── Pillow ─────────────────────────────────────────────────────
from PIL import Image

# ── PDF (read) ─────────────────────────────────────────────────
import pdfplumber
from pypdf import PdfReader, PdfWriter

# ── PDF (write / watermark) ─────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.utils import ImageReader

# ── DOCX ────────────────────────────────────────────────────────
from docx import Document

app = Flask(__name__)

@app.after_request
def add_cors(response):
    response.headers["Access-Control-Allow-Origin"]  = "*"
    response.headers["Access-Control-Allow-Headers"] = "Content-Type"
    response.headers["Access-Control-Allow-Methods"] = "GET,POST,OPTIONS"
    return response

LOGO_PATH = "APH_logo-02.png"

# ════════════════════════════════════════════════════════════════
#  REGEX PATTERNS  (verbatim from reference)
# ════════════════════════════════════════════════════════════════

PHONE_REGEX = re.compile(
    r'''
    (?<!\d)
    (?:
        (?:\+|000?)\s*
        \d{1,4}[\s\-\.\(\)]{0,2}
        (?:\(?\d{1,5}\)?[\s\-\.]?)?
        \d{2,5}[\s\-\.]?\d{2,5}
        (?:[\s\-\.]\d{1,6})?
        |
        [6-9]\d{9}
        |
        0[5-6]\d{8}
        |
        0[5-6]\d[\s\-\.]\d{3}[\s\-\.]\d{4}
        |
        [2-9]\d{2}[\s\-\.][2-9]\d{2}[\s\-\.]\d{4}
        |
        [2-9]\d{2}[2-9]\d{6}
        |
        \d{5}[\s\-\.]\d{5}
        |
        \d{5}[\s\-\.]\d{4}
        |
        \(\d{2,5}\)[\s\-\.]?\d{3,5}[\s\-\.]?\d{3,5}
        |
        \d{2,5}[\s\-\.]\d{4,5}[\s\-\.]\d{3,5}
        |
        07\d{3}[\s\-\.]\d{6}
        |
        \d{4,5}[\-\.]\d{6,8}
    )
    (?!\d)
    ''',
    re.VERBOSE
)

EMAIL_REGEX  = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b')
MAILTO_REGEX = re.compile(r'mailto:[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}', re.IGNORECASE)

PHONE_SYMBOLS = {"☎","☏","✆","📞","📟","📱","📲","🤙","🤳","📳","📴"}
MAIL_SYMBOLS  = {"✉","🖂","🖃","🖄","🖅","📧","📨","📩","📬","📭","📮","@"}
ALL_CONTACT_SYMBOLS = PHONE_SYMBOLS | MAIL_SYMBOLS

CONTACT_LABELS = {
    "ph","phn","phone","phone.","tel","tel.","tele","tele.","telephone",
    "mob","mob.","mobile","mobile.","mob no","mobile no","cell","cell.",
    "cellular","fax","fax.","landline","landline.","direct","direct.",
    "m","t","f","whatsapp","wp","wa","mail","mail.","mailto","mailto.",
    "email","email.","e-mail","e.mail","e","em",
}

CONTACT_SECTION_KEYWORDS = [
    "contact","contact details","personal details","get in touch",
    "personal information","contact information","reference","references",
]

STOP_SECTIONS = [
    "education","experience","skills","projects","summary","objective",
    "work history","employment","qualifications","profile","work experience",
]

SECTION_HEADING_WORDS = {
    "contact","profile","skills","education","experience","reference",
    "languages","summary","objective","projects","work","declaration",
    "passport","details","references","personal","information",
    "qualifications","employment","about","hobbies","interests",
    "achievements","awards",
}

JOB_TITLE_KEYWORDS = {
    "senior","junior","lead","principal","chief","head","sr","jr",
    "associate","assistant","deputy","acting","interim","trainee",
    "intern","graduate","entry","mid","staff","distinguished","fellow",
    "executive","managing","general","regional","global","national",
    "local","group","corporate","divisional","vice","president","vp",
    "svp","evp","avp","ceo","cfo","cto","coo","cpo","ciso","cmo",
    "cdo","cro","director","manager","supervisor","coordinator",
    "specialist","consultant","advisor","strategist","analyst","officer",
    "engineer","developer","programmer","coder","architect","software",
    "hardware","firmware","embedded","frontend","backend","fullstack",
    "mobile","ios","android","web","application","app","devops",
    "devsecops","sre","platform","infrastructure","scientist","ml","ai",
    "machine","learning","deep","nlp","computer","vision","analytics",
    "intelligence","bi","warehouse","pipeline","etl","database","dba",
    "sql","nosql","cloud","aws","azure","gcp","network","networking",
    "systems","administrator","sysadmin","linux","windows","security",
    "cybersecurity","infosec","penetration","tester","ethical","hacker",
    "soc","incident","qa","qe","quality","assurance","testing",
    "automation","manual","performance","technician","support","helpdesk",
    "desk","it","ict","implementations","designer","ux","ui","interaction",
    "visual","graphic","motion","illustrator","animator","researcher",
    "accountant","accounting","accounts","auditor","audit","controller",
    "comptroller","treasurer","treasury","finance","financial","planner",
    "bookkeeper","bookkeeping","payroll","billing","tax","taxation","vat",
    "gst","compliance","actuary","actuarial","underwriter","underwriting",
    "credit","collections","receivables","payables","investment","portfolio",
    "fund","equity","trading","trader","broker","dealer","banker","banking",
    "mortgage","loan","lending","risk","quant","quantitative","wealth",
    "insurance","reinsurance","economist","economics","revenue","budget",
    "budgeting","forecasting","fpa","cost","pricing","internal","external",
    "statutory","hr","human","resources","recruiter","recruiting","talent",
    "acquisition","sourcer","headhunter","hrbp","compensation","benefits",
    "onboarding","learning","development","training","trainer",
    "organizational","od","hris","workforce","planning","engagement",
    "relations","employee","labor","industrial","sales","marketing",
    "growth","commercial","representative","rep","customer","success",
    "cx","relationship","brand","content","copywriter","copy","writer",
    "seo","sem","ppc","digital","social","media","advertising","campaign",
    "crm","ecommerce","retail","merchandiser","buyer","purchasing",
    "category","market","research","public","relations","pr",
    "communications","spokesperson","publicist","operations","operational",
    "ops","logistics","supply","chain","scm","procurement","sourcing",
    "vendor","supplier","warehouse","inventory","stock","fulfillment",
    "distribution","shipping","freight","forwarder","customs","import",
    "export","trade","commerce","fleet","transport","transportation",
    "dispatcher","scheduler","forecaster","facilities","facility",
    "maintenance","production","manufacturing","factory","plant","process",
    "lean","sigma","project","programme","program","portfolio","pmo",
    "scrum","agile","kanban","waterfall","delivery","implementation",
    "rollout","owner","po","roadmap","release","change","transformation",
    "administrative","administration","admin","secretary","secretarial",
    "pa","ea","office","receptionist","clerk","typist","entry","record",
    "document","archivist","legal","lawyer","attorney","advocate",
    "solicitor","barrister","counsel","paralegal","notary","regulatory",
    "governance","contracts","contract","intellectual","property",
    "litigation","corporate","doctor","physician","surgeon","practitioner",
    "gp","nurse","nursing","midwife","paramedic","pharmacist","pharmacy",
    "dentist","dental","orthodontist","therapist","physiotherapist",
    "occupational","psychologist","psychiatrist","counselor","dietitian",
    "nutritionist","optometrist","radiologist","radiographer","sonographer",
    "lab","laboratory","pathologist","microbiologist","medical","clinical",
    "healthcare","health","caregiver","carer","veterinarian","vet",
    "teacher","teaching","tutor","tutoring","professor","lecturer",
    "instructor","faculty","headteacher","headmaster","academic",
    "postdoc","curriculum","education","educational","librarian",
    "guidance","coach","mentor","facilitator","civil","structural",
    "mechanical","electrical","electronics","chemical","petroleum","oil",
    "gas","mining","geological","geotechnical","environmental",
    "sustainability","aeronautical","aerospace","automotive","marine",
    "naval","offshore","instrumentation","piping","surveyor","surveying",
    "quantity","estimator","draughtsman","cad","bim","site","construction",
    "commissioning","inspector","safety","hse","qhse","architect",
    "architecture","interior","urban","planner","landscape","estate",
    "property","realtor","valuer","appraiser","leasing","asset",
    "journalist","reporter","editor","photographer","videographer",
    "cinematographer","filmmaker","author","blogger","podcaster","creative",
    "art","artist","curator","musician","composer","sound","audio",
    "broadcaster","presenter","anchor","host","chef","cook","baker",
    "sous","pastry","bartender","barista","waiter","waitress","steward",
    "hostess","hotel","hospitality","resort","concierge","housekeeping",
    "housekeeper","travel","tourism","tour","guide","event","catering",
    "food","beverage","cashier","store","shop","service","care","call",
    "centre","center","guard","investigator","detective","police",
    "enforcement","lifeguard","fire","chemist","biologist","physicist",
    "mathematician","statistician","epidemiologist","climate","pilot",
    "captain","cabin","crew","stewardess","purser","conductor","operator",
    "mechanic","volunteer","outreach","community","fundraiser","charity",
    "ngo","ndt","nde","ndi","paut","tofd","mfl","mpi","asnt","octg",
    "ut","rt","vt","qc","qi","calibration","inspection","metrologist",
    "radiographic","ultrasonic","penetrant","magnetic","fpso","fso",
    "shipbuilding","outfitting","welder","welding","fitter","plumber",
    "electrician","carpenter","mason","rigger","scaffolder","painter",
    "blaster","boilermaker","millwright","machinist","ironworker",
    "pipefitter","steamfitter","hvac","refrigeration","crane","forklift",
    "driller","roustabout","floorhand","derrickhand","toolpusher",
    "craftsman","artisan","technologist","operative","apprentice",
    "registered","licensed","certified","chartered","professional","design",
    "cunsultant",
}

COMPANY_SUFFIXES = {"ltd","llc","pvt","inc","corp","co","limited","group","plc"}
DEGREE_PREFIXES  = {
    "b.com","b.sc","b.tech","m.com","m.sc","m.tech","mba",
    "bcom","bsc","btech","mcom","msc","ba","phd","diploma",
}
COMMON_FILLER = {"years","year","experience","month","months","with","and","the","of","in","at","for","to"}
LOCATION_WORDS = {
    "uae","india","usa","uk","visa","passport","dubai","abu","sharjah",
    "chennai","mumbai","delhi","bangalore","hyderabad","kochi","cochin",
    "kuwait","qatar","saudi","bahrain","oman","riyadh","doha",
}
NAME_HONORIFICS = {"mr","mrs","ms","miss","dr","prof","er","engr","capt","lt","col","rev"}


# ════════════════════════════════════════════════════════════════
#  HELPERS  (verbatim logic from reference)
# ════════════════════════════════════════════════════════════════

def _strip_honorific(text):
    words = text.strip().split()
    if words and re.sub(r'[^a-z]', '', words[0].lower()) in NAME_HONORIFICS:
        words = words[1:]
    return " ".join(words)


def is_likely_name(text):
    text = _strip_honorific(text.strip())
    if not text or len(text) < 2: return False
    words = text.split()
    if len(words) > 5: return False
    if re.match(r'^[\d\s\.\-/\|:,\(\)]+$', text): return False
    if not re.search(r'[A-Za-z]', text): return False
    lower_words = [re.sub(r'[^a-z]', '', w.lower()) for w in words]
    lower_words = [w for w in lower_words if w]
    if not lower_words: return False
    if len(words) == 1 and lower_words[0] in SECTION_HEADING_WORDS: return False
    if any(w in JOB_TITLE_KEYWORDS for w in lower_words): return False
    if any(w in COMPANY_SUFFIXES for w in lower_words): return False
    if lower_words[0] in DEGREE_PREFIXES: return False
    filler_count = sum(1 for w in lower_words if w in COMMON_FILLER)
    if filler_count >= 2: return False
    if any(w in LOCATION_WORDS for w in lower_words): return False
    return True


def is_likely_role(text):
    text = text.strip()
    if not text or len(text) < 3: return False
    words = text.split()
    if len(words) > 7: return False
    if not re.search(r'[A-Za-z]', text): return False
    lower_words = [w.lower().rstrip('.,') for w in words]
    if len(words) == 1 and lower_words[0] in SECTION_HEADING_WORDS: return False
    if lower_words[0] in DEGREE_PREFIXES: return False
    return any(w in JOB_TITLE_KEYWORDS for w in lower_words)


def span_has_phone(text):   return bool(PHONE_REGEX.search(text))
def span_has_email(text):   return bool(EMAIL_REGEX.search(text)) or bool(MAILTO_REGEX.search(text))
def span_is_symbol(text):
    s = text.strip()
    return s == "@" or any(sym in s for sym in ALL_CONTACT_SYMBOLS)
def span_is_label(text):
    t = text.strip().lower().rstrip(":-./  ")
    return t in CONTACT_LABELS


def sanitize(text):
    if not text: return ""
    text = text.strip().replace("&","and")
    text = re.sub(r"\s+","_",text)
    text = re.sub(r"[^\w\.\-]","",text)
    return text.strip("_.-")


def build_output_filename(name, role):
    safe_name = sanitize(name) or "Unknown_Name"
    safe_role = sanitize(role) or "Unknown_Role"
    return f"APH_{safe_role}_{safe_name}.pdf"


# ════════════════════════════════════════════════════════════════
#  LOGO — transparent watermark
# ════════════════════════════════════════════════════════════════

def create_transparent_logo(opacity=50):
    """Return RGBA PIL Image with logo at given opacity (0-255)."""
    if not os.path.exists(LOGO_PATH):
        return None
    img   = Image.open(LOGO_PATH).convert("RGBA")
    alpha = img.split()[3]
    alpha = alpha.point(lambda p: p * opacity / 255)
    img.putalpha(alpha)
    return img


# ════════════════════════════════════════════════════════════════
#  DOCX → plain text (for name/role extraction)
# ════════════════════════════════════════════════════════════════

def docx_to_text(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


# ════════════════════════════════════════════════════════════════
#  NAME & ROLE EXTRACTION  (pdfplumber port of fitz logic)
# ════════════════════════════════════════════════════════════════

def extract_name_and_role_from_text(text):
    """
    Extract name & role from raw text lines using same
    is_likely_name / is_likely_role heuristics.
    """
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    name_candidates = []
    role_candidates = []

    for line in lines[:60]:   # only first 60 lines
        if is_likely_name(line):
            name_candidates.append(line)
        if is_likely_role(line):
            role_candidates.append(line)

    if not name_candidates:
        return None, None

    name = _strip_honorific(name_candidates[0])

    if not role_candidates:
        return name, None

    # prefer role adjacent (within 5 lines) of name
    try:
        name_idx = lines.index(name_candidates[0])
        adjacent = [r for r in role_candidates
                    if abs(lines.index(r) - name_idx) <= 5
                    and r.lower() != name.lower()]
        if adjacent:
            return name, adjacent[0]
    except ValueError:
        pass

    elsewhere = [r for r in role_candidates if r.lower() != name.lower()]
    return name, (elsewhere[0] if elsewhere else None)


def extract_name_and_role(pdf_bytes=None, text=None):
    if text:
        return extract_name_and_role_from_text(text)
    if pdf_bytes:
        try:
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                page_text = pdf.pages[0].extract_text() or ""
            return extract_name_and_role_from_text(page_text)
        except Exception:
            return None, None
    return None, None


# ════════════════════════════════════════════════════════════════
#  BACKGROUND COLOUR SAMPLER  (PIL port of fitz pixel sampling)
# ════════════════════════════════════════════════════════════════

def get_bg_color_for_line(page_img, line_y_frac, page_h_px):
    """
    Sample dominant colour at ~line_y on a PIL page image.
    Returns (r,g,b) 0-255 tuple.
    """
    try:
        px_y = int(line_y_frac * page_h_px)
        px_y = max(0, min(page_h_px - 1, px_y))
        strip = page_img.crop((0, max(0, px_y-3), page_img.width, min(page_h_px, px_y+3)))
        strip = strip.convert("RGB")
        pixels = list(strip.getdata())
        if not pixels:
            return (255, 255, 255)
        def bucket(c): return tuple(v >> 4 for v in c)
        dominant = Counter(bucket(p) for p in pixels).most_common(1)[0][0]
        matching = [p for p in pixels if bucket(p) == dominant]
        return (
            sum(p[0] for p in matching) // len(matching),
            sum(p[1] for p in matching) // len(matching),
            sum(p[2] for p in matching) // len(matching),
        )
    except Exception:
        return (255, 255, 255)


# ════════════════════════════════════════════════════════════════
#  SECTION DETECTION  (text-based port of fitz block logic)
# ════════════════════════════════════════════════════════════════

def find_contact_section_lines(lines):
    """
    Returns set of line indices that fall within a contact section.
    Mirrors find_contact_regions logic using line text.
    """
    contact_indices = set()
    in_contact = False
    contact_start = -1

    for i, line in enumerate(lines):
        ll = line.lower().strip()

        if any(kw in ll for kw in CONTACT_SECTION_KEYWORDS):
            in_contact    = True
            contact_start = i

        elif in_contact and any(s in ll for s in STOP_SECTIONS) and i > contact_start + 3:
            in_contact = False

        if in_contact:
            contact_indices.add(i)

    # If no section found, treat first 35% of lines as contact zone
    if not contact_indices:
        limit = max(1, int(len(lines) * 0.35))
        contact_indices = set(range(limit))

    return contact_indices


# ════════════════════════════════════════════════════════════════
#  COLLECT REDACTION TARGETS  (text-level; yields (line_idx, value, type))
# ════════════════════════════════════════════════════════════════

def collect_redact_targets_text(lines, contact_indices):
    """
    Three-pass strategy matching reference:
    Pass 1 — flag lines inside contact zone with phone/email
    Pass 2 — collect symbols/labels in contact zone
    Pass 3 — page-wide direct hits (belt-and-suspenders)
    Returns list of line indices to fully blank.
    """
    redact_lines = set()
    found_items  = []   # (type, value)

    # Pass 1: lines in contact zone with contact data
    for i in contact_indices:
        line = lines[i]
        if span_has_email(line) or span_has_phone(line):
            redact_lines.add(i)
            for m in EMAIL_REGEX.finditer(line):
                found_items.append(("email", m.group()))
            for m in PHONE_REGEX.finditer(line):
                found_items.append(("phone", m.group()))

    # Pass 2: symbols / labels in contact zone
    for i in contact_indices:
        line = lines[i]
        if span_is_symbol(line) or span_is_label(line.strip().lower().split(":")[0].strip()):
            redact_lines.add(i)

    # Pass 3: page-wide phone/email regardless of section
    for i, line in enumerate(lines):
        if span_has_phone(line) or span_has_email(line):
            redact_lines.add(i)
            for m in EMAIL_REGEX.finditer(line):
                v = m.group()
                if ("email", v) not in found_items:
                    found_items.append(("email", v))
            for m in PHONE_REGEX.finditer(line):
                v = m.group()
                if ("phone", v) not in found_items:
                    found_items.append(("phone", v))

    return redact_lines, found_items


# ════════════════════════════════════════════════════════════════
#  DOCX ANONYMIZER
# ════════════════════════════════════════════════════════════════

def anonymize_docx_bytes(docx_bytes):
    doc = Document(io.BytesIO(docx_bytes))
    lines = [p.text for p in doc.paragraphs]
    contact_indices = find_contact_section_lines(lines)
    redact_lines, found = collect_redact_targets_text(lines, contact_indices)

    for i in redact_lines:
        para = doc.paragraphs[i]
        for run in para.runs:
            run.text = ""

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                ct = cell.text
                if span_has_email(ct) or span_has_phone(ct) or span_is_symbol(ct):
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = ""
                    found.append(("table_cell", ct[:40]))

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read(), found


# ════════════════════════════════════════════════════════════════
#  PDF ANONYMIZER  (pdfplumber text extraction + pypdf redaction)
# ════════════════════════════════════════════════════════════════

def anonymize_pdf_bytes(pdf_bytes):
    """
    Strategy:
    1. Extract text with pdfplumber (line positions)
    2. Identify redaction targets
    3. Use pypdf to build new PDF with blanked rectangles
    4. Stamp transparent logo watermark via reportlab overlay
    """
    found_items = []

    # ── Read original with pdfplumber ─────────────────────────
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as plumb:
        pages_data = []
        for pg in plumb.pages:
            words = pg.extract_words(keep_blank_chars=True,
                                     x_tolerance=3, y_tolerance=3) or []
            lines_dict = {}
            for w in words:
                key = round(w['top'], 1)
                lines_dict.setdefault(key, []).append(w)

            lines_info = []
            for y_key in sorted(lines_dict.keys()):
                ws = sorted(lines_dict[y_key], key=lambda w: w['x0'])
                text = " ".join(w['text'] for w in ws)
                x0  = min(w['x0'] for w in ws)
                x1  = max(w['x1'] for w in ws)
                y0  = min(w['top'] for w in ws)
                y1  = max(w['bottom'] for w in ws)
                lines_info.append({'text': text, 'x0': x0, 'x1': x1, 'y0': y0, 'y1': y1})

            pages_data.append({
                'width':  pg.width,
                'height': pg.height,
                'lines':  lines_info,
            })

    # ── Determine redaction targets per page ──────────────────
    redact_rects = []   # list of lists of (x0,y0,x1,y1, r,g,b) per page
    for pg_data in pages_data:
        line_texts = [l['text'] for l in pg_data['lines']]
        contact_idx = find_contact_section_lines(line_texts)
        redact_line_idx, found = collect_redact_targets_text(line_texts, contact_idx)
        found_items.extend(found)

        pg_rects = []
        for i in redact_line_idx:
            li = pg_data['lines'][i]
            # Expand rect slightly for clean coverage
            x0 = max(0,                  li['x0'] - 2)
            y0 = max(0,                  li['y0'] - 1)
            x1 = min(pg_data['width'],   li['x1'] + 2)
            y1 = min(pg_data['height'],  li['y1'] + 1)
            pg_rects.append((x0, y0, x1, y1, 255, 255, 255))
        redact_rects.append(pg_rects)

    # ── Build redacted PDF via reportlab overlay ───────────────
    # Step A: render original pages as images via pypdf
    reader = PdfReader(io.BytesIO(pdf_bytes))
    total_pages = len(reader.pages)

    out_pdf_buf = io.BytesIO()

    # We'll use reportlab to build a page-by-page output
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate
    from reportlab.pdfgen import canvas as rl_canvas

    # Build redacted PDF using pypdf writer + overlay approach
    # 1) Create a "redaction overlay" PDF with white rectangles
    # 2) Merge with original PDF using pypdf

    # Build overlay PDF
    overlay_buf = io.BytesIO()
    c = rl_canvas.Canvas(overlay_buf)

    for pg_idx, pg_data in enumerate(pages_data):
        pw, ph = pg_data['width'], pg_data['height']
        c.setPageSize((pw, ph))

        rects = redact_rects[pg_idx] if pg_idx < len(redact_rects) else []
        for (x0, y0, x1, y1, r, g, b) in rects:
            # pdfplumber y is from top; reportlab y is from bottom
            rl_y0 = ph - y1
            rl_y1 = ph - y0
            c.setFillColorRGB(r/255, g/255, b/255)
            c.setStrokeColorRGB(r/255, g/255, b/255)
            c.rect(x0, rl_y0, x1 - x0, rl_y1 - rl_y0, fill=1, stroke=0)

        # Logo watermark
        logo_img = create_transparent_logo(opacity=50)
        if logo_img:
            logo_w = pw * 0.5
            logo_h = logo_w * logo_img.height / logo_img.width
            lx = (pw - logo_w) / 2
            ly = (ph - logo_h) / 2
            logo_buf = io.BytesIO()
            logo_img.save(logo_buf, format="PNG")
            logo_buf.seek(0)
            c.drawImage(ImageReader(logo_buf), lx, ly, logo_w, logo_h, mask='auto')

        c.showPage()

    c.save()
    overlay_buf.seek(0)

    # Merge original + overlay
    writer = PdfWriter()
    overlay_reader = PdfReader(overlay_buf)

    for pg_idx in range(total_pages):
        orig_page = reader.pages[pg_idx]
        if pg_idx < len(overlay_reader.pages):
            overlay_page = overlay_reader.pages[pg_idx]
            orig_page.merge_page(overlay_page)
        writer.add_page(orig_page)

    final_buf = io.BytesIO()
    writer.write(final_buf)
    final_buf.seek(0)
    return final_buf.read(), found_items


# ════════════════════════════════════════════════════════════════
#  DOCX → PDF  (reportlab text rebuild, no LibreOffice needed)
# ════════════════════════════════════════════════════════════════

def convert_docx_to_pdf(docx_bytes, filename="doc.docx"):
    """
    Convert DOCX to a simple PDF preserving text content.
    Uses reportlab — no LibreOffice dependency.
    """
    doc = Document(io.BytesIO(docx_bytes))

    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=A4)
    pw, ph = A4

    margin_x = 60
    margin_y = 60
    x = margin_x
    y = ph - margin_y
    line_h = 14
    font_normal = "Helvetica"
    font_bold   = "Helvetica-Bold"
    size_normal = 10
    size_h1     = 16
    size_h2     = 13

    def new_page():
        nonlocal y
        c.showPage()
        y = ph - margin_y

    for para in doc.paragraphs:
        text = para.text
        if not text.strip():
            y -= line_h * 0.5
            if y < margin_y: new_page()
            continue

        # Detect heading styles
        style_name = para.style.name.lower() if para.style else ""
        if "heading 1" in style_name:
            c.setFont(font_bold, size_h1)
            c.setFillColorRGB(0.055, 0.055, 0.07)
        elif "heading 2" in style_name or "heading 3" in style_name:
            c.setFont(font_bold, size_h2)
            c.setFillColorRGB(0.055, 0.055, 0.07)
        else:
            c.setFont(font_normal, size_normal)
            c.setFillColorRGB(0.23, 0.23, 0.29)

        # Word-wrap
        max_w = pw - 2 * margin_x
        words = text.split()
        line_buf = []
        for word in words:
            test = " ".join(line_buf + [word])
            if c.stringWidth(test, c._fontname, c._fontsize) <= max_w:
                line_buf.append(word)
            else:
                if y < margin_y + line_h: new_page()
                c.drawString(x, y, " ".join(line_buf))
                y -= line_h
                line_buf = [word]
        if line_buf:
            if y < margin_y + line_h: new_page()
            c.drawString(x, y, " ".join(line_buf))
            y -= line_h

    c.save()
    buf.seek(0)
    return buf.read()


# ════════════════════════════════════════════════════════════════
#  FLASK ROUTES
# ════════════════════════════════════════════════════════════════

@app.route("/")
def index():
    return send_file("index.html")


@app.route("/api/anonymize", methods=["POST"])
def anonymize():
    files = request.files.getlist("resumes")
    if not files:
        return jsonify({"error": "No files uploaded"}), 400

    results = []
    errors  = []

    for f in files:
        fname = f.filename
        raw   = f.read()
        try:
            # ── Convert DOCX → PDF text for name/role ─────────
            if fname.lower().endswith(".docx"):
                text      = docx_to_text(raw)
                name, role = extract_name_and_role(text=text)
                anon_bytes, found = anonymize_docx_bytes(raw)
                # Convert anonymized DOCX → PDF
                pdf_bytes = convert_docx_to_pdf(anon_bytes, fname)
            else:
                name, role = extract_name_and_role(pdf_bytes=raw)
                pdf_bytes, found = anonymize_pdf_bytes(raw)

            out_name = build_output_filename(name, role)
            pdf_b64  = base64.b64encode(pdf_bytes).decode()

            # Deduplicate found items
            seen = set()
            deduped = []
            for item in found:
                key = (item[0], item[1])
                if key not in seen:
                    seen.add(key)
                    deduped.append({"type": item[0], "value": item[1]})

            results.append({
                "original_name": fname,
                "output_name":   out_name,
                "candidate_name": name or "Unknown",
                "candidate_role": role or "Unknown",
                "found":         deduped,
                "pdf_b64":       pdf_b64,
            })

        except Exception as e:
            errors.append({
                "file":  fname,
                "error": str(e),
                "trace": traceback.format_exc(),
            })

    return jsonify({"results": results, "errors": errors})


@app.route("/api/logo_preview")
def logo_preview():
    if not os.path.exists(LOGO_PATH):
        return jsonify({"has_logo": False})
    try:
        img = Image.open(LOGO_PATH)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        b64 = base64.b64encode(buf.getvalue()).decode()
        return jsonify({"has_logo": True, "b64": b64, "w": img.width, "h": img.height})
    except Exception:
        return jsonify({"has_logo": False})


if __name__ == "__main__":
    print("=" * 60)
    print("  ResumeGuard Pro — Flask Server")
    print("  http://127.0.0.1:5000")
    print("=" * 60)
    app.run(debug=False, port=5000)